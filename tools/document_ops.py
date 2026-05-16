"""PDF/DOCX операции внутри sandbox."""
from __future__ import annotations
import asyncio
from pathlib import Path

from .sandbox import resolve_in_sandbox, sandbox_root

MAX_PDF_CHARS = 50000  # ограничение чтобы не залить контекст модели


async def read_pdf(path: str) -> str:
    """Извлечь текст из PDF (sandbox-only)."""
    def _do() -> str:
        from pypdf import PdfReader
        target = resolve_in_sandbox(path)
        if not target.exists():
            return f"[read_pdf] '{path}' не существует"
        if not target.is_file():
            return f"[read_pdf] '{path}' не файл"
        try:
            reader = PdfReader(str(target))
        except Exception as e:
            return f"[read_pdf] не удалось открыть PDF: {e}"
        n_pages = len(reader.pages)
        chunks: list[str] = []
        total = 0
        truncated = False
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            chunks.append(f"--- стр. {i+1} ---\n{text.strip()}")
            total += len(text)
            if total > MAX_PDF_CHARS:
                truncated = True
                break
        result = f"PDF: {target.name}, страниц: {n_pages}\n\n" + "\n\n".join(chunks)
        if truncated:
            result += f"\n\n[... обрезано на {MAX_PDF_CHARS} символах, показаны первые {len(chunks)} стр.]"
        return result
    return await asyncio.to_thread(_do)


async def read_docx(path: str) -> str:
    """Извлечь текст из DOCX (параграфы + таблицы)."""
    def _do() -> str:
        from docx import Document
        target = resolve_in_sandbox(path)
        if not target.exists():
            return f"[read_docx] '{path}' не существует"
        if not target.is_file():
            return f"[read_docx] '{path}' не файл"
        try:
            doc = Document(str(target))
        except Exception as e:
            return f"[read_docx] не удалось открыть DOCX: {e}"
        lines: list[str] = [f"DOCX: {target.name}"]
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                style = p.style.name if p.style else ""
                if style.startswith("Heading"):
                    lines.append(f"\n## {txt}")
                else:
                    lines.append(txt)
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n[таблица {ti+1}, {len(table.rows)}x{len(table.columns)}]")
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    return await asyncio.to_thread(_do)


async def write_docx(path: str, title: str = "", content: str = "") -> str:
    """Создать DOCX файл. `content` может содержать строки;
    строки начинающиеся с '# ' / '## ' / '### ' становятся заголовками.
    Файл сохраняется в sandbox (рекомендуется в подпапку output/)."""
    def _do() -> str:
        from docx import Document
        from docx.shared import Pt
        target = resolve_in_sandbox(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        # Базовый стиль
        try:
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
        except Exception:
            pass
        if title:
            doc.add_heading(title, level=0)
        for raw_line in (content or "").splitlines():
            line = raw_line.rstrip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(target))
        rel = target.relative_to(sandbox_root())
        return f"[write_docx] OK, сохранил {target.stat().st_size} байт в {rel}"
    return await asyncio.to_thread(_do)
